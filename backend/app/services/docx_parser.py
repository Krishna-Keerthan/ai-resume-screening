# backend/app/services/docx_parser.py
import docx
import re

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file_path)
        text = ""
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
        
        # Clean the text
        text = clean_text(text)
        return text
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return None

def clean_text(text):
    """Clean extracted text"""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s.,!?@#-]', '', text)
    return text.strip()